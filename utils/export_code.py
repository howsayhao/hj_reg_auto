import docx
import os

doc = docx.Document()

path_to_find = input("Enter the path to find: ")
for root, dirs, files in os.walk(path_to_find):
    for file in files:
        if file.endswith((".py", ".sv", ".jinja", ".jinja2")):
            print(os.path.join(root, file))
            doc.add_paragraph(os.path.join(root, file))
            with open(os.path.join(root, file), "r", encoding="utf-8") as f:
                doc.add_paragraph(f.read())

doc.save("exported_code.docx")



